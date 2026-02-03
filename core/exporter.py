# -*- coding: utf-8 -*-
"""
Modul für den Export von Analyseberichten.

Dieses Modul stellt die Klasse ReportExporter bereit, welche rohen Text
in formatierte Dokumente (DOCX und PDF) umwandelt. Es behandelt dabei
automatisch die Dateipfaderstellung und das Encoding für PDF-Dokumente.

Metadata:
    Author: Diyar Altinses, M.Sc.
    Created: 2026-02-03
"""

from pathlib import Path
from typing import Tuple, Optional

from docx import Document
from fpdf import FPDF


class ReportExporter:
    """
    Verwaltet die Erstellung und Speicherung von Ergebnisberichten.

    Die Klasse bietet Methoden, um aus einem Markdown-ähnlichen Textstring
    sowohl Microsoft Word (.docx) als auch PDF-Dateien zu generieren.
    """

    def save_reports(
        self, 
        text_content: str, 
        original_file_path: str
    ) -> Tuple[Optional[str], Optional[str]]:
        """
        Generiert Berichte basierend auf dem Dateinamen der Eingabedatei.

        Erstellt die Dateien im gleichen Verzeichnis wie die Originaldatei
        und fügt das Suffix "_Report" hinzu.

        Args:
            text_content (str): Der Text, der in die Berichte geschrieben wird.
            original_file_path (str): Der Pfad der analysierten Originaldatei.

        Returns:
            Tuple[Optional[str], Optional[str]]: Ein Tupel bestehend aus
            (Pfad_DOCX, Pfad_PDF). Wenn ein Export fehlschlägt, ist der
            entsprechende Wert None.
        """
        if not original_file_path or not text_content:
            return None, None

        # Nutzung von pathlib für robustes Pfad-Management
        original_path = Path(original_file_path)
        base_name = original_path.stem
        output_dir = original_path.parent
        
        # Basis-Name für den Bericht
        report_base = output_dir / f"{base_name}_Report"

        path_docx = report_base.with_suffix(".docx")
        path_pdf = report_base.with_suffix(".pdf")

        # 1. DOCX erstellen
        success_docx = self._create_docx(text_content, path_docx)

        # 2. PDF erstellen
        success_pdf = self._create_pdf(text_content, path_pdf)

        return (
            str(path_docx) if success_docx else None,
            str(path_pdf) if success_pdf else None
        )

    def _create_docx(self, text: str, path: Path) -> bool:
        """
        Erstellt eine formatierte Word-Datei.

        Interpretiert einfache Markdown-Syntax (#, ##) als Überschriften.

        Args:
            text (str): Der Inhalt.
            path (Path): Der Zielpfad.

        Returns:
            bool: True bei Erfolg, False bei Fehler.
        """
        try:
            doc = Document()
            doc.add_heading('Requirements Analyse Report', 0)

            for line in text.split('\n'):
                line = line.strip()
                if not line:
                    continue
                
                if line.startswith('# '):
                    doc.add_heading(line.replace('# ', ''), level=1)
                elif line.startswith('## '):
                    doc.add_heading(line.replace('## ', ''), level=2)
                elif line.startswith('### '):
                    doc.add_heading(line.replace('### ', ''), level=3)
                else:
                    doc.add_paragraph(line)

            doc.save(path)
            return True

        except Exception as e:
            print(f"Fehler beim Erstellen der DOCX: {e}")
            return False

    def _create_pdf(self, text: str, path: Path) -> bool:
        """
        Erstellt eine PDF-Datei mittels FPDF.

        Da FPDF standardmäßig kein UTF-8 unterstützt, werden Sonderzeichen
        bestmöglich ersetzt oder nach Latin-1 konvertiert.

        Args:
            text (str): Der Inhalt.
            path (Path): Der Zielpfad.

        Returns:
            bool: True bei Erfolg, False bei Fehler.
        """
        try:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)

            # Titel
            pdf.set_font("Arial", 'B', 16)
            pdf.cell(200, 10, txt="Requirements Analyse Report", ln=1, align='C')
            pdf.ln(10)

            # Inhalt zurücksetzen auf Standardgröße
            pdf.set_font("Arial", size=12)

            for line in text.split('\n'):
                # Sanitization für FPDF (Latin-1 Problem)
                clean_line = self._sanitize_for_pdf(line)
                pdf.multi_cell(0, 10, txt=clean_line)

            pdf.output(str(path))
            return True

        except Exception as e:
            print(f"Fehler beim Erstellen der PDF: {e}")
            return False

    def _sanitize_for_pdf(self, text: str) -> str:
        """
        Bereitet Text für FPDF (Latin-1) vor.

        Ersetzt gängige Unicode-Zeichen, die in Latin-1 nicht existieren
        (wie Emojis oder spezielle Bulletpoints), und stellt sicher, dass
        Umlaute korrekt kodiert werden.

        Args:
            text (str): Der Originaltext.

        Returns:
            str: Der kompatible Text string.
        """
        # Mapping für Zeichen, die Latin-1 nicht direkt mag oder
        # die wir spezifisch behandeln wollen.
        replacements = {
            "–": "-",   # En-Dash
            "—": "-",   # Em-Dash
            "„": '"',   # Deutsche Anführungszeichen unten
            "“": '"',   # Deutsche Anführungszeichen oben
            "✔": "[OK]",
            "❌": "[X]",
            "•": "-",   # Bulletpoint
        }
        
        for char, replacement in replacements.items():
            text = text.replace(char, replacement)

        # Versuch, nach Latin-1 zu kodieren. Unbekannte Zeichen werden durch '?' ersetzt.
        return text.encode('latin-1', 'replace').decode('latin-1')